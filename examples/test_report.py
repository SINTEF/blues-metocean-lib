from datetime import datetime

import numpy as np
from bluesmet.common.scatter import Scatter
from bluesmet.normet.nora3 import arome3km_3hr, wave_sub_time
from docx import Document as CreateDoc
from docx.document import Document
#pylint: disable=no-name-in-module
from docx.enum.section import WD_ORIENT
from docx.table import Table
from windrose import WindroseAxes

from bluesmet.normet.metengine.extreme_stats import joint_2D_contour, return_value

# Coordinates we want to get data for
lat_pos = 64.731729
lon_pos = 5.835813
start_date = datetime(1990, 1, 1)
end_date = datetime(2020, 12, 31)
max_concurrent_downloads = 10

document: Document = CreateDoc()

document.add_heading('Metocean', 0)

document.add_paragraph(f'Requested Coordinates: latitude={lat_pos:.4f}, longitude={lon_pos:.4f}')
document.add_paragraph(f'Data for periods: {start_date.strftime("%b/%d/%Y")} to {end_date.strftime("%b/%d/%Y")}')

document.add_page_break()
document.add_heading('Wind data', 1)
# First we present wind data

requested_values=["wind_speed","wind_direction","height"]
years = range(start_date.year,end_date.year)
syears = f"[{years.start}-{years.stop}]"

wind_values = arome3km_3hr.get_values_between(lat_pos,lon_pos,start_date, end_date,requested_values,max_concurrent_downloads=max_concurrent_downloads)

actual_lat = wind_values.latitude
actual_lon = wind_values.longitude

document.add_paragraph(f'Coordinates for collected data: latitude={actual_lat:.4f}, longitude={actual_lon:.4f}')


iheight = 0
speed = wind_values["wind_speed"][:,iheight]
heights = wind_values["height"].values
# Met: North West Up, wind_going_to
# Wind rose: North East Down, wind coming from
direction = np.fmod(wind_values["wind_direction"][:,0]+180.0, 360.0)

ax = WindroseAxes.from_ax()
ax.bar(direction, speed,bins=9,nsector=36, opening=0.8, edgecolor="white")
ax.set_legend()
ax.set_title(f"Wind - {syears} at {heights[iheight]} m")

rose_image_path = "output/wind_rose.png"
ax.figure.savefig(rose_image_path)
width = document.sections[-1].page_width * 0.75
document.add_picture(rose_image_path, width=width)

# Now get the wave data
document.add_page_break()
document.add_heading('Wave data', 1)

requested_values = [
    "hs",
    "tp"
]

return_periods=[5,10,50,100]

# Get the requested data for the given time at the given coordinates
values = wave_sub_time.get_values_between(lat_pos, lon_pos, start_date, end_date, requested_values,max_concurrent_downloads=max_concurrent_downloads)


hs = values["hs"]
actual_lat = wind_values.latitude
actual_lon = wind_values.longitude


document.add_paragraph(f'Coordinates for collected data: latitude={actual_lat:.4f}, longitude={actual_lon:.4f}')

document.add_heading('Basic statistics', level=1)

# Add basic statistics for each requested value
table: Table = document.add_table(rows=1, cols=7)
table.style = 'Table Grid'  # Apply a table style with gridlines
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Value'
hdr_cells[1].text = 'Mean'
hdr_cells[2].text = 'Minimum'
hdr_cells[3].text = 'Maximum'
hdr_cells[4].text = 'p10'
hdr_cells[5].text = 'p90'
hdr_cells[6].text = 'Std'

for requested_value in requested_values:
    tpv = values[requested_value].values
    row_cells = table.add_row().cells
    row_cells[0].text = requested_value
    row_cells[1].text = f"{tpv.mean():.2f}"
    row_cells[2].text = f"{tpv.min():.2f}"
    row_cells[3].text = f"{tpv.max():.2f}"
    row_cells[4].text = f"{np.percentile(tpv, 10):.2f}"
    row_cells[5].text = f"{np.percentile(tpv, 90):.2f}"
    row_cells[6].text = f"{tpv.std():.2f}"

document.add_heading('Contours', level=1)

hs = values["hs"]
tp = values["tp"]

image_path = "output/joint_contour.png"

# Hs - Tp probability contours, 0 - 360 degrees
joint_2D_contour(
    hs, tp, return_periods=[5, 10, 50, 100], image_path=image_path, image_title="Contour"
)

document.add_picture(image_path)

document.add_heading('Return values', level=1)

table: Table = document.add_table(rows=1, cols=2)

table.style = 'Table Grid'  # Apply a table style with gridlines
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Return period (years)'
hdr_cells[1].text = 'Hs (m)'

ret = return_value(hs, return_periods)

for rp,tpv in zip(return_periods,ret.values):
    row_cells = table.add_row().cells
    row_cells[0].text = str(rp)
    row_cells[1].text = f"{tpv:.2f}"


document.add_page_break()
# Shift page orientation to landscape
section = document.add_section()
new_width, new_height = section.page_height, section.page_width
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = new_width
section.page_height = new_height


document.add_heading('Scatter', level=1)

bin_size = 2.0
scatter = Scatter(bin_size=bin_size)
for hsv, tpv in zip(hs.values, tp.values):
    scatter.add(hsv, tpv)

row_vals = scatter.row_values()
col_vals = scatter.column_values()
occurences = scatter.occurences()

# Write the scatter to a word table
table: Table = document.add_table(rows=len(row_vals) + 2, cols=len(col_vals) + 2)
table.style = 'Table Grid'  # Apply a table style with gridlines
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Hs/Tp'
hdr_cells[-1].text = 'Sum'
for i, tpv in enumerate(col_vals):
    hdr_cells[i + 1].text = f"{tpv:.1f}"

for i, occ in enumerate(occurences):
    row_cells = table.rows[i+1].cells
    hsv = row_vals[i]
    row_cells[0].text = f"{hsv:.1f}"
    row = [row_vals[i]] + occ.tolist() + [int(occ.sum())]
    for j, val in enumerate(row):
        if val > 0:
            row_cells[j].text = str(val)

total_sum = occurences.sum()
footer = ["Sum"] + occurences.sum(axis=0).tolist() + [total_sum]
footer_cells = table.rows[-1].cells
for i, val in enumerate(footer):
    footer_cells[i].text = str(val)

document.save('output/metocean.docx')